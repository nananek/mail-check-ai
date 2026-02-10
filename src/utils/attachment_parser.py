import fitz  # PyMuPDF
import io
import csv
import logging
from typing import List, Dict, Tuple, Optional
from pathlib import Path
from PIL import Image

logger = logging.getLogger(__name__)


class AttachmentParser:
    """各種添付ファイルからテキストを抽出するクラス"""
    
    @staticmethod
    def _try_ocr(image_bytes: bytes) -> Optional[str]:
        """画像からOCRでテキストを抽出"""
        try:
            import pytesseract
            from PIL import Image
            
            image = Image.open(io.BytesIO(image_bytes))
            
            # ページの向きを自動検出して回転
            try:
                osd = pytesseract.image_to_osd(image)
                rotation = int([line for line in osd.split('\n') if 'Rotate' in line][0].split(':')[1].strip())
                if rotation != 0:
                    logger.info(f"Detected rotation: {rotation} degrees, rotating image")
                    image = image.rotate(-rotation, expand=True)
            except Exception as e:
                logger.debug(f"Could not detect orientation, using image as-is: {e}")
            
            # OCR実行（日本語+英語）
            text = pytesseract.image_to_string(image, lang='jpn+eng')
            return text.strip()
        
        except ImportError:
            logger.warning("pytesseract not installed, cannot perform OCR")
            return None
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return None
    
    @staticmethod
    def extract_pdf(pdf_bytes: bytes, use_ocr: bool = True) -> str:
        """
        PDFからテキスト抽出（OCR対応）
        
        Args:
            pdf_bytes: PDFファイルのバイナリデータ
            use_ocr: テキストが少ない場合にOCRを使用するか
        """
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            text_parts = []
            total_text_length = 0
            page_images = []  # OCR用に画像を保存
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # まず通常のテキスト抽出を試行
                text = page.get_text()
                text_length = len(text.strip())
                total_text_length += text_length
                
                if text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                    page_images.append(None)  # テキストがあればOCR不要
                else:
                    # テキストがない場合は画像として保存（後でOCR）
                    page_images.append(page_num)
            
            doc_text = "\n\n".join([t for t in text_parts if t])
            
            # テキストがほとんどない（画像PDFの可能性）
            if use_ocr and total_text_length < 100 and page_images:
                logger.info(f"PDF has minimal text ({total_text_length} chars), attempting OCR on {len([p for p in page_images if p is not None])} pages")
                
                ocr_parts = []
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")  # 再度開く
                
                for page_num in range(len(doc)):
                    if page_num in page_images:
                        page = doc.load_page(page_num)
                        
                        # ページを画像に変換（高解像度でOCR精度向上）
                        mat = fitz.Matrix(2.0, 2.0)  # 2倍にスケール
                        pix = page.get_pixmap(matrix=mat)
                        img_bytes = pix.tobytes("png")
                        
                        # OCR実行
                        ocr_text = AttachmentParser._try_ocr(img_bytes)
                        if ocr_text:
                            ocr_parts.append(f"--- Page {page_num + 1} (OCR) ---\n{ocr_text}")
                
                doc.close()
                
                if ocr_parts:
                    # OCRで取得したテキストを追加
                    if doc_text:
                        return doc_text + "\n\n=== OCRによる追加テキスト ===\n\n" + "\n\n".join(ocr_parts)
                    else:
                        return "\n\n".join(ocr_parts)
            
            doc.close()
            return doc_text if doc_text else "[PDFにテキストが見つかりません（画像のみの可能性）]"
        
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            return f"[PDF解析エラー: {str(e)}]"
    
    @staticmethod
    def extract_word(docx_bytes: bytes) -> str:
        """Word文書からテキスト抽出"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(docx_bytes))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # テーブルも抽出
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts) if text_parts else "[Word文書にテキストが見つかりません]"
        
        except ImportError:
            logger.warning("python-docx not installed, cannot parse Word documents")
            return "[Word文書の解析には python-docx が必要です]"
        except Exception as e:
            logger.error(f"Failed to extract Word text: {e}")
            return f"[Word解析エラー: {str(e)}]"
    
    @staticmethod
    def extract_excel(excel_bytes: bytes, max_rows: int = 500) -> str:
        """Excel/CSVからテキスト抽出"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(excel_bytes), read_only=True, data_only=True)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"\n=== シート: {sheet_name} ===")
                
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    if row_count >= max_rows:
                        text_parts.append(f"[... {max_rows}行以降は省略 ...]")
                        break
                    
                    # 空行をスキップ
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(val.strip() for val in row_values):
                        text_parts.append(" | ".join(row_values))
                        row_count += 1
            
            wb.close()
            return "\n".join(text_parts) if text_parts else "[Excelにデータが見つかりません]"
        
        except ImportError:
            logger.warning("openpyxl not installed, cannot parse Excel documents")
            return "[Excel文書の解析には openpyxl が必要です]"
        except Exception as e:
            logger.error(f"Failed to extract Excel text: {e}")
            return f"[Excel解析エラー: {str(e)}]"
    
    @staticmethod
    def extract_csv(csv_bytes: bytes, max_rows: int = 500) -> str:
        """CSVからテキスト抽出"""
        try:
            # 文字コード検出を試みる
            encodings = ['utf-8', 'shift-jis', 'cp932', 'euc-jp', 'iso-2022-jp']
            text = None
            
            for encoding in encodings:
                try:
                    text = csv_bytes.decode(encoding)
                    break
                except (UnicodeDecodeError, LookupError):
                    continue
            
            if text is None:
                text = csv_bytes.decode('utf-8', errors='ignore')
            
            # CSVパース
            csv_reader = csv.reader(io.StringIO(text))
            text_parts = []
            row_count = 0
            
            for row in csv_reader:
                if row_count >= max_rows:
                    text_parts.append(f"[... {max_rows}行以降は省略 ...]")
                    break
                
                if any(cell.strip() for cell in row):
                    text_parts.append(" | ".join(row))
                    row_count += 1
            
            return "\n".join(text_parts) if text_parts else "[CSVにデータが見つかりません]"
        
        except Exception as e:
            logger.error(f"Failed to extract CSV text: {e}")
            return f"[CSV解析エラー: {str(e)}]"
    
    @staticmethod
    def extract_text_file(text_bytes: bytes) -> str:
        """テキストファイルから内容を抽出"""
        try:
            encodings = ['utf-8', 'shift-jis', 'cp932', 'euc-jp', 'iso-2022-jp', 'latin-1']
            
            for encoding in encodings:
                try:
                    return text_bytes.decode(encoding)
                except (UnicodeDecodeError, LookupError):
                    continue
            
            # すべて失敗した場合はエラーを無視してデコード
            return text_bytes.decode('utf-8', errors='ignore')
        
        except Exception as e:
            logger.error(f"Failed to extract text file: {e}")
            return f"[テキストファイル解析エラー: {str(e)}]"
    
    @classmethod
    def extract_from_attachment(cls, filename: str, file_bytes: bytes) -> str:
        """
        ファイル名に基づいて適切な抽出方法を選択
        
        Args:
            filename: ファイル名
            file_bytes: ファイルのバイナリデータ
            
        Returns:
            抽出されたテキスト
        """
        ext = Path(filename).suffix.lower()
        
        # PDF
        if ext == '.pdf':
            return cls.extract_pdf(file_bytes)
        
        # Word
        elif ext in ['.docx', '.doc']:
            if ext == '.docx':
                return cls.extract_word(file_bytes)
            else:
                return "[.doc形式は非対応です。.docx形式に変換してください]"
        
        # Excel
        elif ext in ['.xlsx', '.xls']:
            if ext == '.xlsx':
                return cls.extract_excel(file_bytes)
            else:
                return "[.xls形式は非対応です。.xlsx形式に変換してください]"
        
        # CSV
        elif ext == '.csv':
            return cls.extract_csv(file_bytes)
        
        # テキストファイル
        elif ext in ['.txt', '.log', '.md', '.json', '.xml', '.html', '.htm']:
            return cls.extract_text_file(file_bytes)
        
        # 未対応の形式
        else:
            return f"[未対応のファイル形式: {ext}]"
    
    @classmethod
    def extract_from_multiple(cls, attachments: List[Tuple[str, bytes]], max_size_per_file: int = 10000) -> Dict[str, str]:
        """
        複数の添付ファイルからテキストを抽出
        
        Args:
            attachments: [(filename, file_bytes), ...] のリスト
            max_size_per_file: 1ファイルあたりの最大文字数
            
        Returns:
            {filename: extracted_text} の辞書
        """
        results = {}
        
        for filename, file_bytes in attachments:
            logger.info(f"Extracting text from: {filename}")
            extracted_text = cls.extract_from_attachment(filename, file_bytes)
            
            # 長すぎる場合は切り詰める
            if len(extracted_text) > max_size_per_file:
                extracted_text = extracted_text[:max_size_per_file] + f"\n\n[... 以降 {len(extracted_text) - max_size_per_file} 文字省略 ...]"
            
            results[filename] = extracted_text
        
        return results
